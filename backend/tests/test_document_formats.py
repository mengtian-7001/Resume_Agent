from __future__ import annotations

import json
import shutil
import subprocess
import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document

from app.document_text import (
    DOCX_MIME,
    DOC_MIME,
    PDF_MIME,
    detect_document_mime,
    extract_document_text,
)


class DocumentFormatTests(unittest.TestCase):
    def test_browser_text_builder_produces_parseable_docx(self) -> None:
        node = shutil.which("node")
        if not node:
            self.skipTest("Node.js is not installed")

        root = Path(__file__).resolve().parents[2]
        module_url = (root / "frontend-document.js").as_uri()
        with tempfile.TemporaryDirectory(prefix="resume-agent-text-jd-") as temp_dir:
            output = Path(temp_dir) / "manual-jd.docx"
            source = (
                "岗位名称：AI Agent 工程师\n"
                "岗位职责：负责企业级智能体的设计、开发与上线。\n"
                "任职要求：熟悉 Python、LangChain 和 Function Calling。"
            )
            script = (
                "import { writeFileSync } from 'node:fs';"
                f"import {{ createTextDocxBytes }} from {json.dumps(module_url)};"
                f"writeFileSync({json.dumps(str(output))}, createTextDocxBytes({json.dumps(source)}));"
            )
            completed = subprocess.run(
                [node, "--input-type=module", "-e", script],
                cwd=root,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=15,
                check=False,
            )
            self.assertEqual(completed.returncode, 0, completed.stderr.decode(errors="replace"))
            text = extract_document_text(output.read_bytes(), DOCX_MIME)
            self.assertIn("AI Agent 工程师", text)
            self.assertIn("Function Calling", text)

    def test_scanned_pdf_uses_offline_ocr(self) -> None:
        source = fitz.open()
        source_page = source.new_page(width=800, height=300)
        source_page.insert_text(
            (40, 100),
            "Python FastAPI Agent Engineer 5 years",
            fontsize=32,
        )
        image = source_page.get_pixmap(dpi=180, alpha=False).tobytes("png")
        source.close()

        scanned = fitz.open()
        page = scanned.new_page(width=800, height=300)
        page.insert_image(page.rect, stream=image)
        raw = scanned.tobytes()
        scanned.close()

        text = extract_document_text(raw, PDF_MIME)
        self.assertIn("Python", text)
        self.assertIn("FastAPI", text)

    def test_real_docx_still_extracts_text(self) -> None:
        path = Path(__file__).resolve().parents[2] / "samples" / "resume-cv_001.docx"
        raw = path.read_bytes()
        self.assertEqual(detect_document_mime(raw), DOCX_MIME)
        self.assertIn("林知远", extract_document_text(raw, DOCX_MIME))

    def test_legacy_doc_is_detected_and_converted(self) -> None:
        office = shutil.which("soffice") or shutil.which("libreoffice")
        textutil = shutil.which("textutil")
        if not office and not textutil:
            self.skipTest("DOC converter is not installed")

        with tempfile.TemporaryDirectory(prefix="resume-agent-doc-test-") as temp_dir:
            root = Path(temp_dir)
            docx_path = root / "legacy-source.docx"
            document = Document()
            document.add_heading("AI Agent 工程师", level=1)
            document.add_paragraph("候选人熟悉 Python、FastAPI，并有五年项目经验。")
            document.save(docx_path)

            doc_path = root / "legacy-source.doc"
            if textutil:
                completed = subprocess.run(
                    [textutil, "-convert", "doc", "-output", str(doc_path), str(docx_path)],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=45,
                    check=False,
                )
            else:
                completed = subprocess.run(
                    [office, "--headless", "--convert-to", "doc", "--outdir", str(root), str(docx_path)],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=45,
                    check=False,
                )
            self.assertEqual(completed.returncode, 0, completed.stderr.decode(errors="replace"))
            self.assertTrue(doc_path.exists())

            raw = doc_path.read_bytes()
            self.assertEqual(detect_document_mime(raw), DOC_MIME)
            text = extract_document_text(raw, DOC_MIME)
            self.assertIn("Python", text)
            self.assertIn("FastAPI", text)


if __name__ == "__main__":
    unittest.main()
