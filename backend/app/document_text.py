"""Document text extraction helpers (PDF / DOCX) with magic-byte validation."""

from __future__ import annotations

import os
import shutil
import subprocess
import threading
import zipfile
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"

_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_MAX_DOCUMENT_BYTES = 10 * 1024 * 1024
_MAX_PDF_PAGES = 60
_MAX_OCR_PAGES = 30
_MIN_PAGE_TEXT_CHARS = 24
_MAX_PDF_EDGE_POINTS = 2000
_DOC_CONVERT_TIMEOUT_SEC = 45
_OCR_LOCK = threading.Lock()

# Reject pathological DOCX (zip bombs / oversized members) before python-docx parses.
_MAX_DOCX_UNCOMPRESSED = 40 * 1024 * 1024
_MAX_DOCX_MEMBER = 12 * 1024 * 1024
_MAX_DOCX_FILES = 4000


def detect_document_mime(raw: bytes) -> str | None:
    """Return canonical MIME from file signatures, or None if unrecognized."""
    if len(raw) >= 5 and raw[:5] == b"%PDF-":
        return PDF_MIME
    if len(raw) >= 4 and raw[:4] == b"PK\x03\x04":
        try:
            with zipfile.ZipFile(BytesIO(raw)) as zf:
                names = set(zf.namelist())
                if "[Content_Types].xml" in names and any(
                    name.startswith("word/") for name in names
                ):
                    return DOCX_MIME
        except zipfile.BadZipFile:
            return None
    if len(raw) >= len(_OLE_MAGIC) and raw[: len(_OLE_MAGIC)] == _OLE_MAGIC:
        try:
            import olefile
        except ImportError:
            # The lightweight Vercel API delegates document parsing to the
            # Docker Worker, where the full parser dependency set is installed.
            return None
        try:
            with olefile.OleFileIO(BytesIO(raw)) as ole:
                if ole.exists("WordDocument"):
                    return DOC_MIME
        except (OSError, IOError, olefile.OleFileError):
            return None
    return None


def assert_safe_document(raw: bytes, declared_mime: str | None = None) -> str:
    """
    Validate bytes against magic / ZIP structure and optional declared MIME.
    Returns the detected MIME. Raises ValueError on mismatch or unsafe archives.
    """
    if len(raw) > _MAX_DOCUMENT_BYTES:
        raise ValueError("文件超过 10MB 限制。")
    detected = detect_document_mime(raw)
    if detected is None:
        raise ValueError("无法识别文件类型，请上传真实的 PDF 或 DOCX。")

    declared = (declared_mime or "").split(";")[0].strip().lower()
    if declared and declared not in {detected, "application/octet-stream"}:
        # Allow empty/octet-stream from storage; reject clear mismatches.
        aliases = {
            PDF_MIME: {"application/pdf", "application/x-pdf"},
            DOCX_MIME: {
                DOCX_MIME,
                "application/msword",
                "application/zip",
            },
            DOC_MIME: {DOC_MIME, "application/x-msword", "application/octet-stream"},
        }
        if declared not in aliases.get(detected, {detected}):
            raise ValueError(
                f"文件内容与声明类型不符（声明={declared or '空'}，实际={detected}）。"
            )

    if detected == DOCX_MIME:
        _assert_docx_zip_bounds(raw)
    return detected


def extract_document_text(raw: bytes, mime_type: str) -> str:
    detected = assert_safe_document(raw, mime_type)
    if detected == PDF_MIME:
        return _extract_pdf_text(raw)
    if detected == DOCX_MIME:
        return _extract_docx_text(raw)
    if detected == DOC_MIME:
        return _extract_legacy_doc_text(raw)
    raise ValueError("不支持的文件格式")


def _extract_pdf_text(raw: bytes) -> str:
    """Extract selectable PDF text and OCR only pages with no useful text layer."""
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("PDF 解析组件未安装，请使用独立 Worker。") from exc

    pdf = fitz.open(stream=raw, filetype="pdf")
    try:
        if pdf.page_count > _MAX_PDF_PAGES:
            raise ValueError(f"PDF 页数超过 {_MAX_PDF_PAGES} 页限制。")
        chunks: list[str] = []
        ocr_pages = 0
        for page in pdf:
            if page.rect.width > _MAX_PDF_EDGE_POINTS or page.rect.height > _MAX_PDF_EDGE_POINTS:
                raise ValueError("PDF 页面尺寸异常，已拒绝解析。")
            text = (page.get_text("text") or "").strip()
            meaningful = sum(1 for char in text if char.isalnum())
            if meaningful < _MIN_PAGE_TEXT_CHARS:
                ocr_pages += 1
                if ocr_pages > _MAX_OCR_PAGES:
                    raise ValueError(f"需要 OCR 的页面超过 {_MAX_OCR_PAGES} 页限制。")
                ocr_text = _ocr_pdf_page(page)
                if ocr_text.strip():
                    text = "\n".join(part for part in (text, ocr_text) if part.strip())
            if text:
                chunks.append(text)
        return "\n".join(chunks)
    finally:
        pdf.close()


@lru_cache(maxsize=1)
def _ocr_engine():
    try:
        from rapidocr import RapidOCR
    except ImportError as exc:  # pragma: no cover - guarded by locked requirements
        raise ValueError("OCR 组件未安装，无法解析扫描版 PDF。") from exc
    return RapidOCR()


def _ocr_pdf_page(page: Any) -> str:
    """Render one page at 200 DPI and run offline Chinese/English OCR."""
    try:
        image = page.get_pixmap(dpi=200, alpha=False).tobytes("png")
        # ONNX sessions are cached across tasks; serialize calls for predictable
        # memory use when resume parsing fans out across candidates.
        with _OCR_LOCK:
            result = _ocr_engine()(image)
        texts = getattr(result, "txts", None) or []
        return "\n".join(str(item).strip() for item in texts if str(item).strip())
    except ValueError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"扫描版 PDF OCR 失败：{str(exc)[:160]}") from exc


def _extract_legacy_doc_text(raw: bytes) -> str:
    """Convert a validated Word 97-2003 document to plain text.

    LibreOffice is used on Linux/Docker; macOS can fall back to the built-in
    ``textutil`` command. The conversion happens in an isolated temporary
    directory and is bounded by a hard timeout.
    """
    with TemporaryDirectory(prefix="resume-agent-doc-") as temp_dir:
        root = Path(temp_dir)
        source = root / "input.doc"
        source.write_bytes(raw)

        office = shutil.which("soffice") or shutil.which("libreoffice")
        if office:
            profile = root / "lo-profile"
            profile.mkdir()
            command = [
                office,
                "--headless",
                "--nologo",
                "--nodefault",
                "--nolockcheck",
                "--norestore",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(root),
                str(source),
            ]
            completed = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=_DOC_CONVERT_TIMEOUT_SEC,
                check=False,
            )
            output = root / "input.txt"
            if completed.returncode == 0 and output.exists():
                return _decode_converted_text(output.read_bytes())

        textutil = shutil.which("textutil")
        if textutil:
            completed = subprocess.run(
                [textutil, "-convert", "txt", "-stdout", str(source)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=_DOC_CONVERT_TIMEOUT_SEC,
                check=False,
            )
            if completed.returncode == 0 and completed.stdout:
                return _decode_converted_text(completed.stdout)

    raise ValueError("旧版 DOC 转换器不可用；请在 Worker 安装 LibreOffice Writer。")


def _decode_converted_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "gb18030"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").strip()


def _assert_docx_zip_bounds(raw: bytes) -> None:
    try:
        with zipfile.ZipFile(BytesIO(raw)) as zf:
            infos = zf.infolist()
            if len(infos) > _MAX_DOCX_FILES:
                raise ValueError("DOCX 内文件数量过多，已拒绝解析。")
            total = 0
            for info in infos:
                if info.file_size > _MAX_DOCX_MEMBER:
                    raise ValueError("DOCX 内含过大条目，已拒绝解析。")
                total += max(0, int(info.file_size))
                if total > _MAX_DOCX_UNCOMPRESSED:
                    raise ValueError("DOCX 解压体积过大，已拒绝解析。")
                # Compression ratio sanity for zip bombs.
                if info.compress_size > 0 and info.file_size / max(1, info.compress_size) > 200:
                    raise ValueError("DOCX 压缩比异常，已拒绝解析。")
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX 不是有效的 ZIP 包。") from exc


def _extract_docx_text(raw: bytes) -> str:
    """Read body paragraphs + tables (incl. nested), plus header/footer text."""
    try:
        from docx import Document as DocxDocument
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise ValueError("DOCX 解析组件未安装，请使用独立 Worker。") from exc

    document = DocxDocument(BytesIO(raw))
    chunks: list[str] = []

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = (block.text or "").strip()
            if text:
                chunks.append(text)
        elif isinstance(block, Table):
            table_text = _table_text(block)
            if table_text:
                chunks.append(table_text)

    for section in document.sections:
        for part_name in ("header", "footer"):
            part = getattr(section, part_name, None)
            if part is None:
                continue
            for paragraph in part.paragraphs:
                text = (paragraph.text or "").strip()
                if text:
                    chunks.append(text)
            for table in part.tables:
                table_text = _table_text(table)
                if table_text:
                    chunks.append(table_text)

    return "\n".join(chunks)


def _iter_block_items(parent):
    """Yield paragraphs and tables in document order (python-docx recipe)."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _table_text(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            nested = []
            for nested_table in cell.tables:
                nested_text = _table_text(nested_table)
                if nested_text:
                    nested.append(nested_text)
            cell_text = " ".join(
                (p.text or "").strip() for p in cell.paragraphs if (p.text or "").strip()
            )
            if nested:
                cell_text = " ".join(filter(None, [cell_text, *nested]))
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)
