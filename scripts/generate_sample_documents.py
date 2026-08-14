#!/usr/bin/env python3
"""Build the public sample resumes and keep their web previews in sync."""

from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DIR = ROOT / "samples"
PUBLIC_SAMPLE_DIR = ROOT / "public" / "samples"


PROFILES = [
    {
        "id": "resume-cv_001",
        "name": "林知远",
        "role": "AI Agent / LLM 应用工程师",
        "tag": "推荐",
        "salary": "30-45K",
        "years": 5,
        "summary": "5 年 Python 开发经验，近 4 年专注 LLM 应用与 Agent 生产落地。擅长把复杂业务流程拆成可观测、可回滚的工具调用链路，能够同时承担方案设计、工程实现和上线评测。",
        "education": ["东湖理工大学｜计算机科学与技术｜本科"],
        "experience": [
            {
                "title": "澄空智能｜Agent 平台工程师｜2021.09-至今",
                "bullets": [
                    "主导工单处理 Multi-Agent 的架构与落地，设计规划器、执行器、质检器和人工接管节点，覆盖售后、权限申请与异常排查等 6 类流程。",
                    "使用 LangChain 与自研编排层接入 20+ 内部 API，统一 Function Calling 协议中的鉴权、参数校验、幂等键、超时重试和降级策略。",
                    "建立 480 条核心任务评测集及每日回归机制，将工具误调用率从 11% 降至 3.2%，关键流程一次完成率提升至 86%。",
                    "补齐 OpenTelemetry tracing、调用成本看板和失败样本归因，线上问题平均定位时间由 45 分钟缩短到 12 分钟。",
                    "与产品、客服和安全团队共同定义高风险写操作的审批边界，推动 Agent 从 30 人试点扩展到 600+ 名内部用户。",
                ],
            },
            {
                "title": "拾光纪科技｜Python 后端工程师｜2020.07-2021.03",
                "bullets": [
                    "负责客服中台的工单、用户画像与审计日志接口，使用 FastAPI、PostgreSQL 和 Redis 支撑日均约 8 万次请求。",
                    "梳理 12 个内部系统的权限与错误码规范，为后续将业务能力封装成 Agent 工具打下基础。",
                ],
            },
        ],
        "projects": [
            {
                "title": "“澄助手”企业内部 Copilot｜技术负责人",
                "bullets": [
                    "支持查库、提单、改权限和生成排障报告，日均调用约 1.2 万次，月度活跃率稳定在 71%。",
                    "实现规划—执行—反思循环与并行工具调用；当置信度低或写操作风险高时自动转人工审批。",
                    "将提示词、工具 Schema 和评测基线纳入版本管理，发布时自动执行离线回归与 5% 灰度验证。",
                ],
            }
        ],
        "skills": [
            "语言与框架：Python、FastAPI、Pydantic、LangChain、LangGraph",
            "Agent 工程：Function Calling、Multi-Agent、Prompt Engineering、MCP、RAG",
            "数据与基础设施：PostgreSQL、Redis、Docker、OpenTelemetry、GitHub Actions",
            "工作方式：需求拆解、技术方案评审、跨团队协作、线上故障复盘",
        ],
    },
    {
        "id": "resume-cv_002",
        "name": "周启明",
        "role": "大模型应用开发工程师",
        "tag": "推荐",
        "salary": "28-42K",
        "years": 4,
        "summary": "4 年后端与大模型应用经验，聚焦客服和运营场景的智能体研发。熟悉用中文业务概念抽象工具协议，并通过状态机、失败补偿与评测体系把智能体原型推进到稳定生产服务。",
        "education": [
            "南港大学｜软件工程｜硕士",
            "江宁学院｜软件工程｜本科",
        ],
        "experience": [
            {
                "title": "北纬实验室｜智能体研发工程师｜2022.07-至今",
                "bullets": [
                    "基于 LangChain 搭建 Multi-Agent 协作，由意图识别、知识检索、工具执行和质检智能体完成任务拆解、记忆、重试与人工接管。",
                    "将订单、物流、退款和会员接口封装为 16 个 Tool Calling 工具，补充签名校验、重放防护、失败补偿和人工确认。",
                    "使用 LCEL 与 LangGraph 构建状态机，支持上下文压缩、断点恢复和按业务风险动态路由模型。",
                    "围绕 Function Calling 维护 Prompt Engineering 系统提示词和少样本模板，设计 320 条场景评测集，使核心意图识别准确率达到 93%。",
                    "负责服务监控与成本优化，通过缓存、模型分级和批量摘要将单会话平均推理成本降低 27%。",
                ],
            },
            {
                "title": "问舟信息｜Python 后端开发｜2021.07-2021.12",
                "bullets": [
                    "使用 Python 和 PostgreSQL 开发内部审批流，覆盖权限申请、费用报销和合同用印。",
                    "沉淀统一的权限校验与操作审计模块，后续被智能体侧复用为可调用工具。",
                ],
            },
        ],
        "projects": [
            {
                "title": "“舟问”客服中台｜核心开发",
                "bullets": [
                    "系统日均会话 8000+，自动解决率提升 21%，人工接管率下降 18%。",
                    "为退款、改址等高风险操作设计二次确认与补偿事务，连续 6 个月未发生越权写入事故。",
                    "建设会话 tracing 与质检回放页，支持按工具、模型、提示词版本定位失败原因。",
                ],
            }
        ],
        "skills": [
            "语言与框架：Python、FastAPI、LangChain、LCEL、LangGraph",
            "Agent 工程：Function Calling、Multi-Agent、Prompt Engineering、RAG、状态机编排",
            "工程能力：PostgreSQL、Redis、Docker、日志追踪、离线评测",
            "语言：英语 CET-6，可阅读英文技术文档并参与跨团队技术评审",
        ],
    },
    {
        "id": "resume-cv_003",
        "name": "韩沐辰",
        "role": "LLM 应用 / Agent 工程师",
        "tag": "推荐",
        "salary": "32-48K",
        "years": 6,
        "summary": "6 年对话产品和智能应用研发经验，近 3 年负责可执行 Agent 产品。具备代码助手、浏览器自动化和知识检索等多类工具的生产实践，重视可观测性、失败降级和用户体验。",
        "education": ["西麓大学｜人工智能｜本科"],
        "experience": [
            {
                "title": "星河对话｜高级应用工程师｜2020.07-至今",
                "bullets": [
                    "将单轮问答机器人升级为可执行工具的 Agent，负责需求抽象、技术选型、核心模块开发与上线复盘。",
                    "使用 Python 与 LangChain 实现代码 Agent、浏览器 Agent 双通道，生产环境日活约 3000，峰值每分钟 420 次请求。",
                    "Function Calling 层对接 IDE 插件、内部知识检索和日志平台，统一并发限流、权限票据与参数脱敏。",
                    "Multi-Agent 采用规划器、执行器和评论器分工，失败任务可自动降级为只读建议或转交人工。",
                    "维护 40+ Prompt 模板及自动路由策略，通过结构化输出校验将 JSON 解析失败率降至 0.8%。",
                    "带领 3 人小组完成季度迭代，建立代码评审、灰度发布和线上周度复盘机制。",
                ],
            }
        ],
        "projects": [
            {
                "title": "“星河码助手”研发效能 Agent｜技术负责人",
                "bullets": [
                    "覆盖代码补全、修复单测和日志排障，平均每个任务调用工具 6.4 次，建议采纳率约 64%。",
                    "自研 tracing 记录每步工具参数、模型输出和耗时，支持按失败类型自动聚类复盘。",
                    "为仓库写操作增加沙箱、Diff 确认和回滚点，避免 Agent 直接覆盖用户代码。",
                ],
            },
            {
                "title": "MCP 工具目录试点｜方案设计",
                "bullets": [
                    "调研并接入 5 个 MCP Server，验证工具发现、权限隔离与跨会话复用能力；目前处于小范围试用阶段。",
                ],
            },
        ],
        "skills": [
            "核心技术：Python、LangChain、Function Calling、Multi-Agent、Prompt Engineering",
            "平台能力：FastAPI、SQL、工具编排、并发限流、权限票据、OpenTelemetry",
            "应用方向：代码 Agent、浏览器 Agent、企业知识库、MCP（试点）",
            "团队协作：小组带领、技术评审、项目排期、跨部门需求澄清",
        ],
    },
    {
        "id": "resume-cv_004",
        "name": "陈思齐",
        "role": "大模型应用工程师",
        "tag": "复核",
        "salary": "20-32K",
        "years": 3,
        "summary": "3 年 NLP 应用经验，主要负责企业知识库问答与 FAQ 机器人。具备 Python、LangChain 和基础 Function Calling 实践，但复杂工具编排、高风险写操作和多智能体生产经验仍需在面试中进一步验证。",
        "education": ["青城学院｜计算机科学与技术｜本科"],
        "experience": [
            {
                "title": "木兰计算｜NLP 应用工程师｜2023.07-至今",
                "bullets": [
                    "负责企业知识库问答和 FAQ 机器人，完成文档切分、向量检索、召回重排与引用溯源链路。",
                    "使用 Python、FastAPI 和 LangChain 维护检索服务，服务 12 个客户空间，累计导入约 18 万份文档。",
                    "优化系统提示词、拒答策略和引用格式，使人工抽检准确率由 64% 提升到 71%。",
                    "对接“查文档”“查工单状态”两个只读 Function Calling 工具，具备参数 Schema 和基础异常处理经验。",
                    "参与规划器 Demo 的技术预研，完成任务拆解与顺序执行，但尚未进入正式生产环境。",
                ],
            },
            {
                "title": "木兰计算｜后端开发实习生｜2021.07-2021.10",
                "bullets": [
                    "参与数据标注平台的任务分发和结果导出模块，主要使用 Flask 与 MySQL。",
                ],
            },
        ],
        "projects": [
            {
                "title": "“木兰问问”RAG 知识助手｜核心开发",
                "bullets": [
                    "搭建混合检索与文档引用能力，支持 PDF、Word 和网页内容统一入库。",
                    "建立 200 条问答验收集并跟踪错误类型，当前综合准确率 71%，复杂跨文档问题仍依赖人工。",
                    "完成工单状态查询工具试点；因缺少写操作权限，尚未验证幂等、补偿和审批机制。",
                ],
            }
        ],
        "skills": [
            "熟练：Python、FastAPI、LangChain、RAG、向量检索、Prompt Engineering",
            "了解：Function Calling、Multi-Agent、LangGraph、模型评测",
            "数据组件：PostgreSQL、MySQL、Elasticsearch、Milvus",
            "待加强：复杂任务编排、生产级工具权限、失败重试与可观测性",
        ],
    },
    {
        "id": "resume-cv_005",
        "name": "吴晓岚",
        "role": "LLM 应用开发工程师",
        "tag": "复核",
        "salary": "22-35K",
        "years": 4,
        "summary": "4 年算法和模型评测经验，正在从训练侧转向 LLM 应用工程。具备扎实的 Python、实验设计和数据分析能力，已完成 LangChain 应用 Demo，但生产级工具编排与服务治理经验相对不足。",
        "education": ["滨江大学｜电子信息工程｜本科"],
        "experience": [
            {
                "title": "蓝桉网络｜算法工程师｜2021.07-2025.06",
                "bullets": [
                    "负责文本分类、意图识别和小规模指令微调，覆盖数据清洗、训练、离线评测与人工抽检。",
                    "维护 6 个业务分类模型和统一评测脚本，使用 Precision、Recall、F1 及分层抽样分析版本差异。",
                    "参与垂类客服 SFT 项目，整理 3.5 万条训练样本并设计标注规范，实验集准确率达到 78%。",
                    "2024 年起使用 LangChain 将微调模型封装为对话 Demo，支持检索、摘要和标准话术生成。",
                    "与产品共同维护提示词模板和评测记录，但尚未负责面向真实用户的 SLA、限流与故障恢复。",
                ],
            }
        ],
        "projects": [
            {
                "title": "客服辅助 Copilot 原型｜独立开发",
                "bullets": [
                    "基于 FastAPI、LangChain 和向量库实现问答与回复建议，供销售演示和内部需求评审。",
                    "本地验证天气查询与客户资料检索 Function Calling，完成参数校验和基础超时处理。",
                    "跟随开源教程实现两个角色顺序协作的 Multi-Agent 实验，尚无生产流量和复杂恢复机制。",
                ],
            },
            {
                "title": "模型评测看板｜主要开发",
                "bullets": [
                    "整合离线指标、人工抽检和错误样本，缩短算法版本评审时间约 30%。",
                ],
            },
        ],
        "skills": [
            "熟练：Python、PyTorch、Pandas、模型训练与评测、数据分析",
            "实践：LangChain、FastAPI、RAG、Prompt Engineering",
            "了解：Function Calling、Multi-Agent、Docker、Redis",
            "转岗方向：希望补齐应用服务化、工具权限、评测回归和生产运维能力",
        ],
    },
    {
        "id": "resume-cv_006",
        "name": "赵予安",
        "role": "Agent 开发工程师",
        "tag": "复核",
        "salary": "20-30K",
        "years": 3,
        "summary": "3 年全栈开发经验，熟悉业务后台、接口联调和交付流程，近一年开始参与大模型应用。能够独立完成小型 LangChain 应用，但 Agent 相关深度、评测体系和生产规模仍有限。",
        "education": ["禾山大学｜软件工程｜硕士"],
        "experience": [
            {
                "title": "橙湾科技｜全栈开发工程师｜2023.07-至今",
                "bullets": [
                    "负责项目管理与经营分析后台，使用 Vue、Python 和 PostgreSQL 完成需求开发与持续迭代。",
                    "维护 30+ 内部接口和定时任务，参与权限管理、操作审计及线上问题排查。",
                    "2024 年下半年加入大模型试点，使用 LangChain 开发团队周报摘要 Bot 和日历查询助手。",
                    "为日历查询工具补充参数校验和超时提示；因安全限制，未接入写操作和跨系统事务。",
                    "配合产品调整提示词语气、摘要长度和固定输出格式，尚未建立系统化自动评测集。",
                ],
            }
        ],
        "projects": [
            {
                "title": "内部周报摘要 Bot｜独立开发",
                "bullets": [
                    "聚合 Jira 导出数据与成员周报，生成项目进展、风险和下周计划，服务 12 人研发团队。",
                    "实现长文本分段摘要与结果合并，采用人工抽查方式验收，没有正式 tracing 和成本看板。",
                ],
            },
            {
                "title": "多角色任务助手 Demo｜技术预研",
                "bullets": [
                    "通过两个 Chain 顺序完成资料整理与总结，验证了角色提示词，但尚未实现并行、投票或动态任务分派。",
                ],
            },
        ],
        "skills": [
            "熟练：Python、Vue、PostgreSQL、REST API、业务系统交付",
            "实践：LangChain、Prompt Engineering、Function Calling、FastAPI",
            "了解：Multi-Agent、向量检索、Docker、模型评测",
            "待加强：生产级 Agent 架构、工具幂等与重试、可观测性、大规模用户验证",
        ],
    },
    {
        "id": "resume-cv_007",
        "name": "孙博文",
        "role": "AI Agent 初级工程师",
        "tag": "不匹配",
        "salary": "12-20K",
        "years": 1,
        "summary": "计算机本科毕业约 1 年，具备 Python 基础和 LangChain 入门实践。目前工作以配置、Demo 调试和数据整理为主，尚未达到岗位要求的 3 年经验及独立生产交付能力。",
        "education": ["榕城大学｜计算机科学与技术｜本科"],
        "experience": [
            {
                "title": "云帆数字｜AI 应用实习生 / 初级工程师｜2025.07-至今",
                "bullets": [
                    "在导师指导下维护内部知识问答 Demo，主要负责提示词配置、文档导入和基础问题复现。",
                    "使用 Python 与 LangChain 调整检索参数和输出格式，完成若干小范围功能验证。",
                    "跟做天气 API 的 Function Calling 实验，能够描述工具 Schema，但未参与权限、幂等与失败补偿设计。",
                    "整理人工评测表和错误样本，每周输出准确率统计与典型问题清单。",
                    "阅读 Multi-Agent 开源项目并完成本地运行，尚未独立开发或上线复杂协作流程。",
                ],
            }
        ],
        "projects": [
            {
                "title": "基于 LangChain 的课程问答系统｜毕业设计",
                "bullets": [
                    "完成课程资料切分、向量检索和答案引用展示，支持约 2000 页教学资料。",
                    "使用 80 条手工问题做基础测试，未包含外部工具执行和生产部署。",
                ],
            },
            {
                "title": "天气查询 Agent｜学习项目",
                "bullets": [
                    "调用公开天气 API 并返回结构化结果，主要用于学习 Function Calling 的参数定义。",
                ],
            },
        ],
        "skills": [
            "基础：Python、Git、SQL、REST API、Linux 常用命令",
            "入门实践：LangChain、RAG、Prompt Engineering、Function Calling",
            "了解：Multi-Agent、FastAPI、Docker",
            "当前短板：从业年限、独立系统设计、生产上线与故障处理经验",
        ],
    },
    {
        "id": "resume-cv_008",
        "name": "郑小禾",
        "role": "LLM 应用工程师",
        "tag": "不匹配",
        "salary": "15-25K",
        "years": 5,
        "summary": "5 年以上 Python 内部工具开发经验，近 3 年参与销售助手和 CRM 自动化，具备 Function Calling 与小范围 Multi-Agent 试点经验。技术能力接近岗位，但最高学历为大专，不满足当前 JD 的本科硬门槛。",
        "education": [
            "滨海职业技术学院｜计算机应用技术｜大专",
            "继续教育：完成计算机专业专升本课程学习，未取得本科学历证书",
        ],
        "experience": [
            {
                "title": "磐石软件｜应用开发工程师｜2021.08-至今",
                "bullets": [
                    "长期使用 Python 开发销售运营与数据处理工具，覆盖客户导入、商机提醒、报价校验和日报生成。",
                    "2023 年起参与销售助手建设，使用 LangChain 接入 CRM 查询、跟进记录写入和话术检索。",
                    "为 CRM 写入工具增加字段白名单、请求幂等键和失败重试，降低重复跟进记录问题。",
                    "维护按行业和销售阶段分类的提示词库，支持 8 类话术推荐与异议处理场景。",
                    "参与销售、售前两个角色的 Multi-Agent 试点，完成线索摘要到方案建议的顺序协作。",
                    "负责内部用户培训与问题收集，推动试点覆盖 30 人，但项目尚无正式 SLA。",
                ],
            }
        ],
        "projects": [
            {
                "title": "销售 Copilot 与 CRM 回写｜核心开发",
                "bullets": [
                    "整合客户画像、历史沟通和产品资料，生成下一步跟进建议并支持人工确认后回写 CRM。",
                    "通过缓存和提示词压缩将平均响应时间从 8.1 秒降低到 5.6 秒。",
                    "建立 120 条典型销售问题清单，采用人工评分跟踪内容相关性和合规性。",
                ],
            },
            {
                "title": "销售 / 售前协作 Agent 试点｜主要开发",
                "bullets": [
                    "销售角色提炼客户需求，售前角色生成方案大纲；当前为顺序协作，未覆盖复杂动态规划。",
                ],
            },
        ],
        "skills": [
            "熟练：Python、FastAPI、SQL、REST API、业务自动化",
            "实践：LangChain、Function Calling、Multi-Agent、Prompt Engineering、RAG",
            "工程组件：PostgreSQL、Redis、Docker、GitLab CI",
            "客观限制：最高学历为大专，未满足目标岗位的本科及以上硬性要求",
        ],
    },
]


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for text in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(text)


def write_resume_docx(profile: dict, path: Path) -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(profile["name"])
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(17, 17, 17)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(f"求职意向：{profile['role']}　｜　{profile['years']} 年经验　｜　期望：{profile['salary']}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("个人概况", level=1)
    doc.add_paragraph(profile["summary"])

    doc.add_heading("工作经历", level=1)
    for job in profile["experience"]:
        doc.add_heading(job["title"], level=2)
        add_bullets(doc, job["bullets"])

    doc.add_heading("项目经历", level=1)
    for project in profile["projects"]:
        doc.add_heading(project["title"], level=2)
        add_bullets(doc, project["bullets"])

    doc.add_heading("专业技能", level=1)
    add_bullets(doc, profile["skills"])

    doc.add_heading("教育经历", level=1)
    add_bullets(doc, profile["education"])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("简历中台公开演示样例 · 人物、组织与数据均为虚构")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(125, 125, 125)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def preview_text(profile: dict) -> str:
    lines = [
        f"姓名：{profile['name']}",
        f"求职意向：{profile['role']}",
        f"工作年限：{profile['years']} 年",
        f"期望薪资：{profile['salary']}",
        "",
        "【个人概况】",
        profile["summary"],
        "",
        "【工作经历】",
    ]
    for job in profile["experience"]:
        lines.append(job["title"])
        lines.extend(f"- {item}" for item in job["bullets"])
        lines.append("")
    lines.append("【项目经历】")
    for project in profile["projects"]:
        lines.append(project["title"])
        lines.extend(f"- {item}" for item in project["bullets"])
        lines.append("")
    lines.append("【专业技能】")
    lines.extend(f"- {item}" for item in profile["skills"])
    lines.extend(["", "【教育经历】"])
    lines.extend(f"- {item}" for item in profile["education"])
    return "\n".join(lines).strip()


def main() -> None:
    manifest_path = SAMPLE_DIR / "manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["resumes"] = []

    for profile in PROFILES:
        filename = f"{profile['name']} · 简历.docx"
        docx_name = f"{profile['id']}.docx"
        output_path = SAMPLE_DIR / docx_name
        write_resume_docx(profile, output_path)
        manifest["resumes"].append(
            {
                "id": profile["id"],
                "title": profile["name"],
                "filename": filename,
                "path": f"samples/{docx_name}",
                "tag": profile["tag"],
                "salary": profile["salary"],
                "years": profile["years"],
                "preview": preview_text(profile),
            }
        )

    manifest_text = json.dumps(manifest, ensure_ascii=False, indent=2) + "\n"
    manifest_path.write_text(manifest_text, encoding="utf-8")
    PUBLIC_SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    (PUBLIC_SAMPLE_DIR / "manifest.json").write_text(manifest_text, encoding="utf-8")
    for profile in PROFILES:
        docx_name = f"{profile['id']}.docx"
        shutil.copyfile(SAMPLE_DIR / docx_name, PUBLIC_SAMPLE_DIR / docx_name)

    print(f"Updated {len(PROFILES)} rich sample resumes in samples/ and public/samples/.")


if __name__ == "__main__":
    main()
